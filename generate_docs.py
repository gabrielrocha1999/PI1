"""
Executa este script para gerar a documentação do projeto MindTrack em .docx.
Requer: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)
    return heading


def add_paragraph(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    return table


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.style = "No Spacing"
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F7FF")
    pPr.append(shd)
    return p


BLUE = (46, 111, 181)
GREEN = (67, 160, 71)


def main():
    doc = Document()

    # ── Title Page ────────────────────────────────────────────────
    title = doc.add_heading("MindTrack", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(*BLUE)
        run.font.size = Pt(32)

    sub = doc.add_paragraph("Documentação Técnica do Sistema")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(90, 120, 140)

    date_p = doc.add_paragraph(f"Gerado em: {datetime.date.today().strftime('%d/%m/%Y')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(10)
    date_p.runs[0].italic = True

    doc.add_page_break()

    # ── 1. Visão Geral ────────────────────────────────────────────
    add_heading(doc, "1. Visão Geral do Sistema", 1, BLUE)
    add_paragraph(doc,
        "O MindTrack é uma aplicação web full stack desenvolvida para apoiar pacientes em "
        "acompanhamento psicológico baseado na Terapia Cognitivo-Comportamental (TCC). "
        "O sistema permite que pacientes registrem atividades do cotidiano, marquem como "
        "concluídas e avaliem sua satisfação com notas de 1 a 5 estrelas. Psicólogos podem "
        "acompanhar o progresso de seus pacientes por meio de um painel dedicado."
    )
    add_paragraph(doc, "Objetivos principais:")
    for item in [
        "Facilitar o registro de atividades terapêuticas do dia a dia;",
        "Permitir avaliação de satisfação e reflexão ao concluir tarefas;",
        "Oferecer ao psicólogo uma visão consolidada do progresso do paciente;",
        "Garantir segurança dos dados com autenticação JWT e senhas com hash bcrypt.",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # ── 2. Arquitetura ────────────────────────────────────────────
    add_heading(doc, "2. Arquitetura do Sistema", 1, BLUE)
    add_paragraph(doc,
        "O sistema adota uma arquitetura cliente-servidor com separação clara entre "
        "frontend (React) e backend (FastAPI). A comunicação é feita exclusivamente via "
        "API REST com JSON sobre HTTP."
    )

    add_heading(doc, "2.1 Frontend", 2, GREEN)
    add_table(doc,
        ["Tecnologia", "Versão", "Finalidade"],
        [
            ["React", "18.x", "Biblioteca de interface"],
            ["Vite", "5.x", "Bundler e servidor de desenvolvimento"],
            ["React Router", "6.x", "Roteamento SPA"],
            ["Axios", "1.x", "Cliente HTTP"],
            ["CSS Puro", "—", "Estilização com variáveis CSS"],
        ]
    )

    add_heading(doc, "2.2 Backend", 2, GREEN)
    add_table(doc,
        ["Tecnologia", "Versão", "Finalidade"],
        [
            ["Python", "3.11+", "Linguagem base"],
            ["FastAPI", "0.110", "Framework web REST"],
            ["SQLAlchemy", "2.0", "ORM"],
            ["SQLite", "—", "Banco de dados"],
            ["python-jose", "3.3", "Geração e verificação JWT"],
            ["passlib[bcrypt]", "1.7", "Hash de senhas"],
            ["Pydantic v2", "2.6", "Validação de dados"],
        ]
    )

    add_heading(doc, "2.3 Estrutura de Pastas (Backend)", 2, GREEN)
    add_code_block(doc, """backend/
├── app/
│   ├── main.py               # App FastAPI + CORS + roteadores
│   ├── database.py           # Engine e SessionLocal SQLAlchemy
│   ├── models/models.py      # Modelos ORM: User, Task
│   ├── schemas/schemas.py    # Schemas Pydantic de entrada/saída
│   ├── routers/
│   │   ├── deps.py           # Dependências (autenticação JWT)
│   │   ├── auth.py           # POST /auth/register, POST /auth/login
│   │   ├── tasks.py          # GET/POST/PUT/DELETE /tasks/
│   │   └── psychologist.py   # GET /psychologist/patients/...
│   └── services/
│       └── auth_service.py   # hash_password, verify_password, JWT
├── .env
└── requirements.txt""")

    doc.add_paragraph()

    # ── 3. Modelagem de Entidades ─────────────────────────────────
    add_heading(doc, "3. Modelagem das Entidades", 1, BLUE)

    add_heading(doc, "3.1 Tabela: users", 2, GREEN)
    add_table(doc,
        ["Campo", "Tipo", "Descrição"],
        [
            ["id", "INTEGER (PK)", "Identificador único autoincremental"],
            ["name", "VARCHAR(150)", "Nome completo do usuário"],
            ["email", "VARCHAR(255) UNIQUE", "E-mail de acesso (único)"],
            ["senha_hash", "VARCHAR(255)", "Hash bcrypt da senha"],
            ["tipo", "VARCHAR(20)", "'patient' ou 'psychologist'"],
        ]
    )

    add_heading(doc, "3.2 Tabela: tasks", 2, GREEN)
    add_table(doc,
        ["Campo", "Tipo", "Descrição"],
        [
            ["id", "INTEGER (PK)", "Identificador único autoincremental"],
            ["titulo", "VARCHAR(200)", "Título da tarefa"],
            ["descricao", "TEXT", "Descrição detalhada (opcional)"],
            ["data_prevista", "VARCHAR(20)", "Data prevista no formato YYYY-MM-DD"],
            ["concluida", "BOOLEAN", "Indica se a tarefa foi concluída"],
            ["satisfacao", "INTEGER", "Nota de 1 a 5 (preenchida na conclusão)"],
            ["reflexao", "TEXT", "Reflexão do paciente ao concluir (opcional)"],
            ["user_id", "INTEGER (FK → users.id)", "Paciente dono da tarefa"],
        ]
    )

    doc.add_paragraph()

    # ── 4. Documentação da API ────────────────────────────────────
    add_heading(doc, "4. Documentação da API REST", 1, BLUE)
    add_paragraph(doc, "Base URL: http://localhost:8000")
    add_paragraph(doc, "Rotas protegidas requerem o header: Authorization: Bearer <token>")

    add_heading(doc, "4.1 Autenticação", 2, GREEN)

    add_heading(doc, "POST /auth/register", 3)
    add_paragraph(doc, "Cadastra um novo usuário no sistema.", italic=True)
    add_paragraph(doc, "Body (JSON):", bold=True)
    add_code_block(doc, '{ "name": "string", "email": "string", "senha": "string", "tipo": "patient|psychologist" }')
    add_paragraph(doc, "Resposta 201:", bold=True)
    add_code_block(doc, '{ "id": 1, "name": "João Silva", "email": "joao@email.com", "tipo": "patient" }')

    add_heading(doc, "POST /auth/login", 3)
    add_paragraph(doc, "Autentica um usuário e retorna o token JWT.", italic=True)
    add_paragraph(doc, "Body (JSON):", bold=True)
    add_code_block(doc, '{ "email": "string", "senha": "string" }')
    add_paragraph(doc, "Resposta 200:", bold=True)
    add_code_block(doc, '{ "access_token": "eyJ...", "token_type": "bearer", "user": { "id": 1, "name": "...", "email": "...", "tipo": "patient" } }')

    add_heading(doc, "4.2 Tarefas (Paciente)", 2, GREEN)
    add_table(doc,
        ["Método", "Endpoint", "Descrição"],
        [
            ["GET", "/tasks/", "Lista todas as tarefas do paciente autenticado"],
            ["POST", "/tasks/", "Cria nova tarefa"],
            ["GET", "/tasks/{id}", "Retorna uma tarefa específica"],
            ["PUT", "/tasks/{id}", "Atualiza tarefa (inclusive conclusão e satisfação)"],
            ["DELETE", "/tasks/{id}", "Remove uma tarefa"],
        ]
    )

    add_heading(doc, "Body PUT /tasks/{id} (conclusão):", 3)
    add_code_block(doc, '{ "concluida": true, "satisfacao": 4, "reflexao": "Texto opcional..." }')

    add_heading(doc, "4.3 Psicólogo", 2, GREEN)
    add_table(doc,
        ["Método", "Endpoint", "Descrição"],
        [
            ["GET", "/psychologist/patients", "Lista todos os pacientes cadastrados"],
            ["GET", "/psychologist/patients/{id}", "Retorna dados de um paciente"],
            ["GET", "/psychologist/patients/{id}/tasks", "Lista tarefas de um paciente"],
        ]
    )

    doc.add_paragraph()

    # ── 5. Fluxo de Autenticação JWT ──────────────────────────────
    add_heading(doc, "5. Fluxo de Autenticação com JWT", 1, BLUE)
    steps = [
        ("1. Cadastro", "O usuário envia nome, e-mail, senha e tipo. A senha é processada com bcrypt e armazenada como hash. Nenhuma senha é salva em texto claro."),
        ("2. Login", "O usuário envia e-mail e senha. O backend verifica o hash bcrypt. Se válido, gera um JWT com sub=user_id e tipo=patient|psychologist, com expiração de 24h."),
        ("3. Requisições autenticadas", "O frontend armazena o token no localStorage e o envia no header Authorization: Bearer <token> em todas as requisições protegidas."),
        ("4. Validação pelo backend", "O middleware de dependência (deps.py) decodifica o JWT em cada requisição. Se inválido ou expirado, retorna 401. Se o tipo não corresponde à rota, retorna 403."),
        ("5. Logout", "O frontend remove o token e os dados do usuário do localStorage, encerrando a sessão local."),
    ]
    for title_step, description in steps:
        p = doc.add_paragraph(style="List Number")
        run_title = p.add_run(title_step + ": ")
        run_title.bold = True
        p.add_run(description)

    doc.add_paragraph()

    # ── 6. Instruções de Execução ─────────────────────────────────
    add_heading(doc, "6. Instruções de Execução Local", 1, BLUE)

    add_heading(doc, "6.1 Backend", 2, GREEN)
    add_code_block(doc, """cd backend
python -m venv venv
venv\\Scripts\\activate        # Windows
source venv/bin/activate      # Linux/Mac
pip install -r requirements.txt
uvicorn app.main:app --reload""")
    add_paragraph(doc, "API disponível em: http://localhost:8000")
    add_paragraph(doc, "Swagger UI: http://localhost:8000/docs")

    add_heading(doc, "6.2 Frontend", 2, GREEN)
    add_code_block(doc, """cd frontend
npm install
npm run dev""")
    add_paragraph(doc, "Frontend disponível em: http://localhost:5173")

    add_heading(doc, "6.3 Gerar esta documentação", 2, GREEN)
    add_code_block(doc, """pip install python-docx
python generate_docs.py""")

    doc.add_paragraph()

    # ── 7. Segurança ──────────────────────────────────────────────
    add_heading(doc, "7. Considerações de Segurança", 1, BLUE)
    for item in [
        "Senhas armazenadas exclusivamente como hash bcrypt (fator de custo padrão: 12).",
        "Tokens JWT assinados com chave secreta configurável via variável de ambiente.",
        "Autorização granular: pacientes só acessam suas próprias tarefas; psicólogos acessam dados de pacientes (somente leitura).",
        "CORS configurado para aceitar apenas origens conhecidas (localhost:5173 e localhost:3000).",
        "Validação de entrada com Pydantic v2 em todas as rotas, incluindo validação do campo satisfacao (1–5).",
        "Interceptor Axios redireciona para /login em respostas 401, removendo tokens expirados.",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    # Save
    output = "Documentacao_MindTrack.docx"
    doc.save(output)
    print(f"Documentacao gerada: {output}")


if __name__ == "__main__":
    main()
